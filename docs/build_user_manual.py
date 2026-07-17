from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "manual_assets"
OUT = ROOT / "docs" / "SpiritLens_新手使用手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)
    paragraph.add_run(" 页")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("SpiritLens 新手使用手册")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(20, 37, 70)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("适合第一次使用 SpiritLens 的创作者、运营人员和管理员")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(5.0)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("适用版本", "当前本地项目：Next.js 前端 + FastAPI 后端，访问路径包含 /spiritlens"),
        ("推荐读者", "首次登录、第一次创作图片/视频、管理项目资产或查看后台的用户"),
        ("阅读方式", "先完成“启动与登录”，再按需要阅读图片、视频、画布、项目和后台章节"),
        ("截图说明", "截图取自本地演示环境；若后端未启动，页面可能出现模型或数据加载失败提示"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_shading(table.rows[i].cells[0], "E8EEF5")
        set_cell_text(table.rows[i].cells[0], k, bold=True, color=DARK_BLUE)
        set_cell_text(table.rows[i].cells[1], v)

    doc.add_paragraph()


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = DARK_BLUE
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)


def add_figure(doc, image_name, caption):
    image_path = ASSETS / image_name
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_quick_table(doc):
    doc.add_heading("1. 快速认识 SpiritLens", level=1)
    doc.add_paragraph(
        "SpiritLens 是一个面向 AI 创意制作的工作台，核心能力包括 AI 图片生成、AI 视频生成、智能画布、资产库、灵感社区、影视项目工作流和管理员后台。"
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["入口", "主要用途", "第一次使用建议"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    data = [
        ("我的工作台", "查看个人作品、进入资产库或新建创作", "登录后先看这里，确认账号状态正常"),
        ("AI 图片", "输入提示词生成图片，可上传参考图并设置尺寸、批量数等参数", "先用一段简单提示词测试模型是否可用"),
        ("AI 视频", "根据文字或图片生成视频，适合短片和镜头素材", "先确认模型列表和额度配置"),
        ("智能画布", "用节点串联文本、图片、视频和上传素材", "适合复杂创意流程，先从模板或双击添加节点开始"),
        ("项目", "管理剧集、脚本、角色、场景、道具和分镜", "需要连续内容生产时优先使用"),
        ("灵感社区", "查看和发布作品，进行点赞和评论", "发布前检查作品标题、描述和可见性"),
        ("管理员后台", "管理用户、AI 模型、日志、社区和系统设置", "仅管理员账号可用，生产环境需修改默认密码"),
    ]
    for row_data in data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            set_cell_text(row.cells[i], text)
    doc.add_paragraph()


def build():
    doc = Document()
    style_document(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.text = "SpiritLens 使用手册"
    header.runs[0].font.name = "Microsoft YaHei"
    header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    add_page_number(doc.sections[0].footer.paragraphs[0])

    add_title(doc)
    add_quick_table(doc)

    doc.add_heading("2. 启动与访问", level=1)
    doc.add_heading("2.1 本地启动顺序", level=2)
    add_numbers(
        doc,
        [
            "先启动后端服务，默认地址为 http://localhost:8000，并确认 /docs 接口文档可打开。",
            "再启动前端服务，默认地址为 http://localhost:3000。",
            "本项目配置了 basePath，浏览器访问时使用 http://localhost:3000/spiritlens。",
            "若部署在服务器或反向代理后面，请把前端、后端和静态资源路径统一到同一个域名或代理规则下。",
        ],
    )
    add_callout(
        doc,
        "首次排查重点",
        "如果页面能打开但模型、作品或项目列表加载失败，通常是后端未启动、API 地址配置不一致、Token 失效或数据库/Redis 服务不可用。先检查浏览器控制台和后端日志。",
    )

    doc.add_heading("2.2 登录与注册", level=2)
    doc.add_paragraph("未登录时，受保护页面会提示先登录。普通用户可以进入注册页创建账号，也可以在登录页输入已有账号继续使用。")
    add_figure(doc, "01-login.png", "图 1：登录页，包含普通登录入口和管理员登录入口")
    add_numbers(
        doc,
        [
            "点击右上角“登录”，进入登录页。",
            "输入邮箱或用户名、密码后提交。",
            "登录成功后会进入工作台；右上角会显示当前用户昵称。",
            "如果是管理员，用户菜单中会出现后台入口，也可以从登录页进入管理员登录。",
        ],
    )
    add_figure(doc, "02-register.png", "图 2：注册页，用于创建新账号")

    doc.add_heading("3. 我的工作台", level=1)
    doc.add_paragraph("工作台是普通用户登录后的起点，用来查看近期作品、进入资产库、快速开始图片创作。")
    add_figure(doc, "03-workspace.png", "图 3：工作台空状态，适合第一次创建作品")
    add_bullets(
        doc,
        [
            "“新建创作”会进入 AI 图片生成页。",
            "“资产库”用于查看已经生成或上传的图片、视频素材。",
            "当账号已有作品时，工作台会显示作品统计和最近创作缩略图。",
        ],
    )

    doc.add_heading("4. AI 图片生成", level=1)
    doc.add_paragraph("AI 图片页采用对话式工作流。左侧是对话历史，中间是结果区，底部是提示词输入框和生成参数。")
    add_figure(doc, "04-image.png", "图 4：AI 图片生成页")
    add_numbers(
        doc,
        [
            "在底部输入框描述想要的画面，例如主体、风格、镜头、光线和画幅。",
            "点击模型名称可以切换可用图片模型。",
            "点击参数按钮设置尺寸、批量数量、负向提示词和 seed。",
            "需要参考图时，点击图片按钮上传或从已有作品中选择。",
            "点击发送按钮开始生成；完成后可以预览、下载、复用提示词或发布到社区。",
        ],
    )
    add_callout(
        doc,
        "提示词建议",
        "第一次测试时不要写得过长。建议先写清楚主体和风格，确认模型能正常返回结果后，再逐步增加构图、材质、光照、色彩和负向约束。",
    )

    doc.add_heading("5. AI 视频生成", level=1)
    doc.add_paragraph("AI 视频页与图片页结构相似，但生成目标是视频。它适合把文字描述、首尾帧或参考图片转成动态镜头。")
    add_figure(doc, "05-video.png", "图 5：AI 视频生成页")
    add_bullets(
        doc,
        [
            "先选择视频模型，再填写动作、镜头运动、场景氛围和时长等要求。",
            "图生视频时尽量使用清晰、主体明确的参考图。",
            "生成过程可能比图片更久，可在当前会话中查看进度。",
            "如果长时间没有进度，检查后端任务队列、Redis、Celery worker 和模型供应商配置。",
        ],
    )

    doc.add_heading("6. 智能画布", level=1)
    doc.add_paragraph("智能画布用于把文本、图片、视频和上传节点组合成一个可视化流程，适合需要多步生成或素材复用的创作。")
    add_figure(doc, "06-canvas.png", "图 6：智能画布页")
    add_numbers(
        doc,
        [
            "在空白画布中双击，添加文本、图片、视频或上传节点。",
            "拖动节点边缘调整大小，拖动节点位置整理流程。",
            "把文本节点连接到图片节点，可以把文本作为图片生成提示。",
            "把图片节点连接到视频节点，可以把生成图作为视频参考。",
            "使用缩放、框选和拖拽整理大型工作流。",
        ],
    )

    doc.add_heading("7. 项目与剧集工作流", level=1)
    doc.add_paragraph("项目页适合连续内容制作，例如短剧、动画、分镜和系列视频。项目下可以继续管理季、集、脚本、角色、场景、道具和导演工作台。")
    add_figure(doc, "07-projects.png", "图 7：项目列表页")
    add_bullets(
        doc,
        [
            "点击“新建项目”创建项目，进入项目详情后再创建剧集。",
            "剧集工作台通常按脚本、角色/场景、导演工作台、成片导出和提示词管理组织。",
            "导入小说或脚本后，可以拆解角色、场景、道具和分镜，再继续生成视觉资产。",
            "删除项目会影响其剧集和关联数据，生产环境操作前应确认备份。",
        ],
    )

    doc.add_heading("8. 灵感社区与资产库", level=1)
    doc.add_paragraph("灵感社区用于浏览公开作品、发布自己的生成结果、点赞和评论。资产库则用于管理个人生成素材。")
    add_figure(doc, "08-community.png", "图 8：灵感社区页")
    add_bullets(
        doc,
        [
            "图片或视频生成完成后，可以从结果菜单发布到社区。",
            "发布前建议补充标题、描述和标签，便于其他用户理解作品。",
            "资产库是个人素材管理入口，适合找回历史生成结果和复用参考图。",
        ],
    )

    doc.add_heading("9. 管理员后台", level=1)
    doc.add_paragraph("管理员后台用于平台维护，普通创作者不需要进入。默认管理员账号通常用于本地测试，生产环境必须更换密码并限制访问。")
    add_figure(doc, "09-admin.png", "图 9：管理员后台")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["模块", "用途"]):
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for k, v in [
        ("仪表盘", "查看用户、作品、模型调用和系统运行概况。"),
        ("用户管理", "管理账号状态、权限和用户信息。"),
        ("模型管理", "配置图片、视频、文本模型及成本、尺寸、批量等能力。"),
        ("使用记录", "查看生成任务、调用日志和错误信息。"),
        ("社区管理", "处理公开作品、评论和内容审核。"),
        ("系统设置", "维护平台级配置。"),
    ]:
        row = table.add_row()
        set_cell_text(row.cells[0], k)
        set_cell_text(row.cells[1], v)

    doc.add_heading("10. 常见问题", level=1)
    qa = [
        ("页面打开是 404", "确认访问路径是否包含 /spiritlens。本地默认地址通常是 http://localhost:3000/spiritlens。"),
        ("登录后数据为空", "新账号没有作品属于正常现象；如果项目、模型都加载失败，检查后端 API 地址和数据库连接。"),
        ("模型列表加载失败", "检查后端是否启动，管理员后台是否配置模型，供应商 API Key 是否有效。"),
        ("生成任务一直运行", "检查 Redis、Celery worker、WebSocket 连接和模型供应商接口返回。"),
        ("图片或视频无法预览", "检查文件存储路径、后端静态资源路由、反向代理的上传目录映射。"),
        ("管理员入口看不到", "确认当前用户的 is_admin 权限，或使用管理员专用登录入口。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(q)
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        doc.add_paragraph(a)

    doc.add_heading("11. 第一次使用检查清单", level=1)
    add_bullets(
        doc,
        [
            "前端地址能访问，并且路径包含 /spiritlens。",
            "后端 /docs 能打开，数据库迁移已完成。",
            "注册或登录成功，右上角能看到用户昵称。",
            "管理员后台已配置至少一个可用图片模型和视频模型。",
            "完成一次简单图片生成，并确认结果进入资产库。",
            "需要视频功能时，确认任务队列、Redis 和视频供应商配置正常。",
            "生产部署前修改默认管理员密码，检查 HTTPS、反向代理、上传目录和备份策略。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
